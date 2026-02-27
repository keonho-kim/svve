"""
목적:
- 문서 파일을 파싱해 페이지 노드로 변환하는 전처리 파이프라인을 제공한다.

설명:
- Markdown/PDF/DOCX 입력을 처리한다.
- PDF 표/이미지, DOCX 표를 주석 서비스와 연결해 본문에 반영한다.
- 추출 블록을 max_chunk_chars 기준으로 결합해 page 노드 목록으로 변환한다.

디자인 패턴:
- 파이프라인(Pipeline) + 전략 분기(확장자별 처리).

참조:
- src_py/vtree_search/ingestion/docx_layout.py
- src_py/vtree_search/ingestion/parser_helpers.py
- src_py/vtree_search/contracts/ingestion_models.py
"""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Iterable

import pypdfium2 as pdfium
from PIL import Image

from vtree_search.config.models import IngestionConfig
from vtree_search.contracts.ingestion_models import IngestionPageNode
from vtree_search.exceptions import ConfigurationError, DependencyUnavailableError, IngestionProcessingError
from vtree_search.ingestion.docx_layout import (
    advance_docx_page_state,
    estimate_docx_paragraph_layout,
    estimate_docx_table_height,
    is_docx_page_break_before,
    iterate_docx_blocks,
    resolve_docx_layout_metrics,
)
from vtree_search.ingestion.parser_helpers import (
    chunk_blocks,
    estimate_docx_body_font_size,
    extract_pdf_image_boxes,
    is_usable_image,
    max_docx_font_size,
    pick_one_file_per_extension,
    resolve_docx_heading_level,
    serialize_docx_table,
    table_matrix_to_html,
    to_ltree_label,
    to_pixel_box,
)
from vtree_search.ingestion.source_types import ExtractedBlock
from vtree_search.llm.langchain_ingestion import LangChainIngestionAnnotationLLM

SUPPORTED_SUFFIXES = {".pdf", ".md", ".markdown", ".docx"}
_RENDER_SCALE = 2.0


class SourceParser:
    """파일 전처리/청킹/노드 생성을 담당하는 파서."""

    def __init__(
        self,
        config: IngestionConfig,
        annotation_llm: LangChainIngestionAnnotationLLM | None,
    ) -> None:
        self._config = config
        self._annotation_llm = annotation_llm

    def scan_input_files(self, input_root: str | Path, sample: bool | None = None) -> list[Path]:
        """입력 루트에서 지원 확장자 파일 목록을 수집한다."""
        root = Path(input_root)
        if not root.exists():
            raise IngestionProcessingError(f"입력 경로가 존재하지 않습니다: {root}")

        files = [
            path
            for path in sorted(root.rglob("*"))
            if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES
        ]

        if sample is None:
            sample = self._config.preprocess.sample_per_extension
        if sample:
            files = list(pick_one_file_per_extension(files))
        return files

    async def build_page_nodes_from_files(
        self,
        *,
        document_id: str,
        parent_node_id: str,
        input_root: str | Path,
        sample: bool | None = None,
    ) -> list[IngestionPageNode]:
        """파일 루트에서 페이지 노드 목록을 생성한다."""
        paths = self.scan_input_files(input_root, sample=sample)
        if not paths:
            return []

        extracted: list[ExtractedBlock] = []
        for path in paths:
            extracted.extend(await self._extract_blocks(path))

        chunked = chunk_blocks(extracted, max_chars=self._config.preprocess.max_chunk_chars)
        return _to_page_nodes(
            document_id=document_id,
            parent_node_id=parent_node_id,
            blocks=chunked,
        )

    async def _extract_blocks(self, path: Path) -> list[ExtractedBlock]:
        suffix = path.suffix.lower()
        if suffix in {".md", ".markdown"}:
            return self._extract_markdown(path)
        if suffix == ".docx":
            return await self._extract_docx(path)
        if suffix == ".pdf":
            return await self._extract_pdf(path)
        return []

    def _extract_markdown(self, path: Path) -> list[ExtractedBlock]:
        text = path.read_text(encoding="utf-8")
        paragraphs = [" ".join(part.split()) for part in text.split("\n\n") if part.strip()]

        blocks: list[ExtractedBlock] = []
        for index, paragraph in enumerate(paragraphs, start=1):
            blocks.append(
                ExtractedBlock(
                    source_file=path.as_posix(),
                    page_num=index,
                    block_type="paragraph",
                    text=paragraph,
                    metadata={"layout_type": "markdown", "page_num": index},
                )
            )
        return blocks

    async def _extract_docx(self, path: Path) -> list[ExtractedBlock]:
        try:
            from docx import Document as WordDocument
        except Exception as exc:  # noqa: BLE001
            raise DependencyUnavailableError(
                f"DOCX 파싱을 위해 python-docx가 필요합니다: {exc}"
            ) from exc

        document = WordDocument(str(path))
        paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if getattr(paragraph, "text", None) and str(paragraph.text).strip()
        ]
        body_font_size = estimate_docx_body_font_size(paragraphs)
        layout_metrics = resolve_docx_layout_metrics(document)
        usable_width_pt = float(layout_metrics["usable_width_pt"])
        usable_height_pt = float(layout_metrics["usable_height_pt"])

        blocks: list[ExtractedBlock] = []
        page_texts: dict[int, list[str]] = {}
        paragraph_index = 0
        table_index = 0
        current_page_num = 1
        used_height_pt = 0.0

        for block_type, block in iterate_docx_blocks(document):
            if block_type == "paragraph":
                paragraph = block
                text = " ".join(str(getattr(paragraph, "text", "") or "").split())
                if not text:
                    continue

                if is_docx_page_break_before(paragraph):
                    current_page_num += 1
                    used_height_pt = 0.0

                paragraph_index += 1
                font_size = max_docx_font_size(paragraph) or body_font_size
                style = getattr(paragraph, "style", None)
                style_name = str(getattr(style, "name", "") or "").strip()
                heading_level = resolve_docx_heading_level(
                    style_name=style_name,
                    font_size=font_size,
                    body_font_size=body_font_size,
                    text=text,
                )
                paragraph_layout = estimate_docx_paragraph_layout(
                    paragraph=paragraph,
                    text=text,
                    font_size=font_size,
                    usable_width_pt=usable_width_pt,
                )
                block_height_pt = float(paragraph_layout["estimated_height_pt"])
                page_num, current_page_num, used_height_pt = advance_docx_page_state(
                    current_page_num=current_page_num,
                    used_height_pt=used_height_pt,
                    block_height_pt=block_height_pt,
                    usable_height_pt=usable_height_pt,
                )

                metadata: dict[str, object] = {
                    "layout_type": "docx_paragraph",
                    "page_num": page_num,
                    "block_type": "heading" if heading_level is not None else "paragraph",
                    "heading_tag": f"H{heading_level}" if heading_level is not None else "BODY",
                    "paragraph_index": paragraph_index,
                    "font_size": round(float(font_size), 2),
                    "body_font_size": round(float(body_font_size), 2),
                    "line_spacing": round(float(paragraph_layout["line_spacing_pt"]), 2),
                    "char_spacing": round(float(paragraph_layout["char_spacing_pt"]), 2),
                    "space_before": round(float(paragraph_layout["space_before_pt"]), 2),
                    "space_after": round(float(paragraph_layout["space_after_pt"]), 2),
                    "line_count_estimated": int(paragraph_layout["line_count_estimated"]),
                    "estimated_height_pt": round(block_height_pt, 2),
                    "page_size": "A4",
                }
                if heading_level is not None:
                    metadata["heading_level"] = heading_level

                blocks.append(
                    ExtractedBlock(
                        source_file=path.as_posix(),
                        page_num=page_num,
                        block_type="heading" if heading_level is not None else "paragraph",
                        text=text,
                        metadata=metadata,
                    )
                )
                page_texts.setdefault(page_num, []).append(text)
                continue

            table = block
            html = serialize_docx_table(table)
            if not html:
                continue
            if not self._config.preprocess.enable_table_annotation:
                continue

            table_index += 1
            table_height_pt = estimate_docx_table_height(
                table=table,
                body_font_size=body_font_size,
                usable_width_pt=usable_width_pt,
            )
            page_num, current_page_num, used_height_pt = advance_docx_page_state(
                current_page_num=current_page_num,
                used_height_pt=used_height_pt,
                block_height_pt=table_height_pt,
                usable_height_pt=usable_height_pt,
            )
            annotated = await self._annotate_table(
                table_html=html,
                page_text="\n".join(page_texts.get(page_num, [])),
            )
            blocks.append(
                ExtractedBlock(
                    source_file=path.as_posix(),
                    page_num=page_num,
                    block_type="table",
                    text=annotated,
                    metadata={
                        "layout_type": "docx_table",
                        "page_num": page_num,
                        "block_type": "table",
                        "table_index": table_index,
                        "heading_tag": "BODY",
                        "body_font_size": round(float(body_font_size), 2),
                        "estimated_height_pt": round(float(table_height_pt), 2),
                        "page_size": "A4",
                    },
                )
            )

        return blocks

    async def _extract_pdf(self, path: Path) -> list[ExtractedBlock]:
        table_by_page = await self._extract_pdf_tables(path)
        image_by_page = await self._extract_pdf_images(path)
        pdf = pdfium.PdfDocument(str(path))
        blocks: list[ExtractedBlock] = []
        try:
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                page_num = page_index + 1
                text_page = None
                try:
                    text_page = page.get_textpage()
                    text = " ".join(text_page.get_text_range().split())

                    if text:
                        blocks.append(
                            ExtractedBlock(
                                source_file=path.as_posix(),
                                page_num=page_num,
                                block_type="paragraph",
                                text=text,
                                metadata={"layout_type": "pdf_text", "page_num": page_num},
                            )
                        )

                    for table_index, table_body in enumerate(table_by_page.get(page_num, []), start=1):
                        blocks.append(
                            ExtractedBlock(
                                source_file=path.as_posix(),
                                page_num=page_num,
                                block_type="table",
                                text=table_body,
                                metadata={
                                    "layout_type": "pdf_table",
                                    "page_num": page_num,
                                    "table_index": table_index,
                                },
                            )
                        )

                    for image_index, (image_body, image_path) in enumerate(
                        image_by_page.get(page_num, []),
                        start=1,
                    ):
                        blocks.append(
                            ExtractedBlock(
                                source_file=path.as_posix(),
                                page_num=page_num,
                                block_type="image",
                                text=image_body,
                                metadata={
                                    "layout_type": "pdf_image",
                                    "page_num": page_num,
                                    "image_index": image_index,
                                    "image_path": image_path.as_posix(),
                                },
                            )
                        )
                finally:
                    if text_page is not None:
                        text_page.close()
                    page.close()
        finally:
            pdf.close()
        return blocks

    async def _extract_pdf_tables(self, path: Path) -> dict[int, list[str]]:
        if not self._config.preprocess.enable_table_annotation:
            return {}

        try:
            import pdfplumber
        except Exception as exc:  # noqa: BLE001
            raise DependencyUnavailableError(
                f"PDF 표 추출을 위해 pdfplumber가 필요합니다: {exc}"
            ) from exc

        tables_by_page: dict[int, list[str]] = {}
        with pdfplumber.open(str(path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                page_text = " ".join((page.extract_text() or "").split())
                raw_tables = page.extract_tables() or []
                if not raw_tables:
                    continue

                for raw_table in raw_tables:
                    html = table_matrix_to_html(raw_table)
                    if not html:
                        continue
                    annotated = await self._annotate_table(
                        table_html=html,
                        page_text=page_text,
                    )
                    tables_by_page.setdefault(page_index, []).append(annotated)

        return tables_by_page

    async def _extract_pdf_images(self, path: Path) -> dict[int, list[tuple[str, Path]]]:
        if not self._config.preprocess.enable_image_annotation:
            return {}

        try:
            import pdfplumber
        except Exception as exc:  # noqa: BLE001
            raise DependencyUnavailableError(
                f"PDF 이미지 위치 추출을 위해 pdfplumber가 필요합니다: {exc}"
            ) from exc

        output_dir = Path(self._config.preprocess.asset_output_dir).expanduser().resolve()
        output_dir.mkdir(parents=True, exist_ok=True)
        target_dir = output_dir / path.stem
        target_dir.mkdir(parents=True, exist_ok=True)

        pdf = pdfium.PdfDocument(str(path))
        images_by_page: dict[int, list[tuple[str, Path]]] = {}
        try:
            with pdfplumber.open(str(path)) as plumber_pdf:
                page_count = min(len(pdf), len(plumber_pdf.pages))
                for page_index in range(page_count):
                    page_num = page_index + 1
                    pdfium_page = pdf[page_index]
                    text_page = None
                    rendered = None
                    rendered_image: Image.Image | None = None

                    try:
                        text_page = pdfium_page.get_textpage()
                        page_text = " ".join(text_page.get_text_range().split())
                        plumber_page = plumber_pdf.pages[page_index]
                        image_boxes = extract_pdf_image_boxes(plumber_page)
                        if not image_boxes:
                            continue

                        rendered = pdfium_page.render(scale=_RENDER_SCALE)
                        rendered_image = rendered.to_pil()
                        image_width, image_height = rendered_image.size
                        page_width = float(getattr(plumber_page, "width", 0.0) or 0.0)
                        page_height = float(getattr(plumber_page, "height", 0.0) or 0.0)

                        for image_index, box in enumerate(image_boxes, start=1):
                            pixel_box = to_pixel_box(
                                box=box,
                                page_width=page_width,
                                page_height=page_height,
                                image_width=image_width,
                                image_height=image_height,
                            )
                            if pixel_box is None:
                                continue

                            cropped = rendered_image.crop(pixel_box)
                            try:
                                if not is_usable_image(cropped):
                                    continue
                                image_path = target_dir / f"page-{page_num:04d}-img-{image_index:03d}.png"
                                cropped.save(str(image_path), format="PNG")
                            finally:
                                cropped.close()

                            annotated = await self._annotate_image(
                                image_path=image_path,
                                page_text=page_text,
                            )
                            images_by_page.setdefault(page_num, []).append((annotated, image_path))
                    finally:
                        if rendered_image is not None:
                            rendered_image.close()
                        if rendered is not None:
                            rendered.close()
                        if text_page is not None:
                            text_page.close()
                        pdfium_page.close()
        finally:
            pdf.close()
        return images_by_page

    async def _annotate_table(self, *, table_html: str, page_text: str) -> str:
        annotation = self._require_annotation_llm(kind="표")
        return await annotation.annotate_table(table_html=table_html, page_text=page_text)

    async def _annotate_image(self, *, image_path: Path, page_text: str) -> str:
        annotation = self._require_annotation_llm(kind="이미지")
        return await annotation.annotate_image(image_path=image_path, page_text=page_text)

    def _require_annotation_llm(self, kind: str) -> LangChainIngestionAnnotationLLM:
        if self._annotation_llm is None:
            raise ConfigurationError(
                f"{kind} 주석이 활성화되어 있으나 llm 인자가 없습니다"
            )
        return self._annotation_llm


def build_source_parser(
    config: IngestionConfig,
    annotation_llm: LangChainIngestionAnnotationLLM | None = None,
) -> SourceParser:
    """설정 기반 SourceParser를 생성한다."""
    return SourceParser(config=config, annotation_llm=annotation_llm)


def _to_page_nodes(
    *,
    document_id: str,
    parent_node_id: str,
    blocks: Iterable[ExtractedBlock],
) -> list[IngestionPageNode]:
    doc_label = to_ltree_label(document_id)
    parent_label = to_ltree_label(parent_node_id)

    nodes: list[IngestionPageNode] = []
    for index, block in enumerate(blocks, start=1):
        node_id = uuid.uuid4().hex
        path = f"{doc_label}.{parent_label}.p{max(1, block.page_num)}.b{index}"
        metadata = dict(block.metadata)
        metadata["source_file"] = block.source_file
        metadata["block_type"] = block.block_type

        image_url = str(metadata.get("image_path") or "").strip() or None

        nodes.append(
            IngestionPageNode(
                node_id=node_id,
                parent_node_id=parent_node_id,
                document_id=document_id,
                path=path,
                content=block.text,
                image_url=image_url,
                metadata=metadata,
            )
        )

    return nodes
